"""DOCX generation via python-docx.

Parses markdown structure: headings (h1-h3), paragraphs, fenced code, simple tables.
"""
import re
from pathlib import Path
from typing import Any, Dict

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_code_block(doc: Document, code: str, language: str = ""):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)


def _add_table_from_md(doc: Document, lines):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    header = rows[0]
    body = [r for r in rows[1:] if not all(set(c) <= set("-: ") for c in r)]
    if not body and len(rows) >= 2 and all(set(c) <= set("-: ") for c in "".join(rows[1])):
        body = []
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(body, start=1):
        for ci, val in enumerate(row):
            if ci < len(table.rows[ri].cells):
                table.rows[ri].cells[ci].text = val


def render(
    title: str,
    content: str,
    output_path: Path,
    metadata: Dict[str, Any] = None,
) -> Path:
    metadata = metadata or {}
    doc = Document()

    # Title
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if metadata.get("author"):
        p = doc.add_paragraph()
        r = p.add_run(f"by {metadata['author']}")
        r.italic = True
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # Parse content line by line, handling fenced code blocks and tables
    lines = content.split("\n")
    i = 0
    in_code = False
    code_buf = []
    code_lang = ""
    table_buf = []

    def flush_code():
        nonlocal code_buf, code_lang
        if code_buf:
            _add_code_block(doc, "\n".join(code_buf), code_lang)
        code_buf = []
        code_lang = ""

    def flush_table():
        nonlocal table_buf
        if table_buf:
            _add_table_from_md(doc, table_buf)
        table_buf = []

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
                code_lang = line.strip().strip("`").strip()
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            table_buf.append(line)
            i += 1
            continue
        else:
            flush_table()

        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue
        m = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            level = min(len(m.group(1)), 6)
            doc.add_heading(m.group(2), level=level)
            i += 1
            continue
        if stripped.startswith(("- ", "* ", "+ ")):
            try:
                doc.add_paragraph(stripped[2:], style="List Bullet")
            except KeyError:
                doc.add_paragraph(stripped[2:])
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            try:
                doc.add_paragraph(text, style="List Number")
            except KeyError:
                doc.add_paragraph(text)
            i += 1
            continue
        if stripped == "---":
            doc.add_paragraph("─" * 40)
            i += 1
            continue
        # Plain paragraph (collect continuation lines)
        para = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if not nxt.strip():
                break
            if re.match(r"^(#{1,6}\s|\*|- |\+ |\d+\.\s|\|)", nxt.strip()):
                break
            if nxt.strip().startswith("```"):
                break
            para.append(nxt.strip())
            i += 1
        # Inline markdown: **bold** *italic* `code`
        para_text = " ".join(para)
        para_text = re.sub(r"\*\*(.+?)\*\*", r"\1", para_text)
        para_text = re.sub(r"\*(.+?)\*", r"\1", para_text)
        para_text = re.sub(r"`(.+?)`", r"\1", para_text)
        doc.add_paragraph(para_text)

    flush_code()
    flush_table()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
